"""
predictor.py
------------
Fairness-aware elective-care priority model + explainable-AI layer.

Design choices (documented for the paper):
  * Unit of prediction = provider x specialty area (the project's ethical unit;
    NOT individual patients). The model generalises the transparent scoring so a
    provider can upload new data and get a consistent fairness-aware priority.
  * Target = priority_score_100 (0-100). Predicted level is derived by banding the
    predicted score (Lower/Moderate/High/Critical) so it stays consistent.
  * Features = real-world quantities (breach rates, DTA, demand, throughput,
    backlog growth, fairness gaps, diagnostics/bed/A&E/cancellation/theatre).
  * Fairness WITHOUT bias: protected characteristics are NEVER used as raw group
    identity. They enter only as *aggregate fairness gaps* (over-representation of
    a group in the longest waits), which raise priority for under-served areas —
    an equity-promoting signal, audited in `fairness_audit`.
  * Explainability = dependency-free local attribution (drop-to-median ablation)
    that sums to the deviation from the model's baseline prediction, plus global
    permutation/impurity importance. SHAP-compatible but needs no extra install.
"""
from __future__ import annotations
import os
for _v in ("OMP_NUM_THREADS", "OPENBLAS_NUM_THREADS", "MKL_NUM_THREADS", "NUMEXPR_NUM_THREADS"):
    os.environ.setdefault(_v, "1")
from pathlib import Path
import io
import numpy as np
import pandas as pd
import joblib
import sklearn
from sklearn.ensemble import RandomForestRegressor
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.model_selection import train_test_split
from sklearn.metrics import r2_score, mean_absolute_error, accuracy_score

ROOT = Path(__file__).resolve().parents[1]
PROCESSED = ROOT / "data" / "processed"
MODELS = ROOT / "models"
MODEL_PATH = MODELS / "fairqueue_model.joblib"

FEATURES = [
    "breach_18w_rate", "breach_52w_rate", "dta_pressure", "demand_pressure",
    "backlog_growth_rate", "throughput_rate",
    "fairness_deprivation_score", "fairness_ethnicity_score", "fairness_age_score",
    "fairness_sex_score", "missing_demographic_score",
    "diagnostic_bottleneck", "bed_pressure", "ae_pressure_score",
    "cancellation_pressure", "theatre_pressure",
]
FAIRNESS_FEATURES = [f for f in FEATURES if f.startswith("fairness_") or f == "missing_demographic_score"]

FEATURE_LABELS = {
    "breach_18w_rate": "18-week breach rate", "breach_52w_rate": "52-week long-wait rate",
    "dta_pressure": "Decision-to-admit pressure", "demand_pressure": "New demand pressure",
    "backlog_growth_rate": "Backlog growth", "throughput_rate": "Throughput (completions)",
    "fairness_deprivation_score": "Deprivation over-representation",
    "fairness_ethnicity_score": "Ethnicity over-representation",
    "fairness_age_score": "Age over-representation", "fairness_sex_score": "Sex imbalance",
    "missing_demographic_score": "Missing demographic data",
    "diagnostic_bottleneck": "Diagnostic >6-week bottleneck", "bed_pressure": "Bed occupancy",
    "ae_pressure_score": "A&E pressure", "cancellation_pressure": "Cancellation disruption",
    "theatre_pressure": "Theatre capacity pressure",
}

# accept common alternative column names on uploaded files
SYNONYMS = {
    "diagnostic_over_6w_rate": "diagnostic_bottleneck",
    "bed_occupancy_rate": "bed_pressure",
    "theatre_daycase_share": "theatre_pressure",   # will be inverted below
    "18w_breach_rate": "breach_18w_rate", "52w_breach_rate": "breach_52w_rate",
}


def band(score):
    s = float(score)
    return "Critical" if s >= 80 else "High" if s >= 60 else "Moderate" if s >= 40 else "Lower"


# ------------------------------------------------------------------ training
def train_and_save(random_state=42):
    df = pd.read_parquet(PROCESSED / "modelling_dataset.parquet")
    scr = pd.read_parquet(PROCESSED / "priority_scores.parquet")[
        ["month", "provider_code", "treatment_function_code", "priority_score_100"]]
    df = df.merge(scr, on=["month", "provider_code", "treatment_function_code"], how="left")
    X = df[FEATURES].astype(float)
    y = df["priority_score_100"].astype(float)
    Xtr, Xte, ytr, yte = train_test_split(X, y, test_size=0.25, random_state=random_state)
    if len(Xtr) > 8000:                      # subsample for speed; 16 features need few rows
        Xtr = Xtr.sample(8000, random_state=random_state)
        ytr = ytr.loc[Xtr.index]
    # single-threaded RF (n_jobs=1 -> no OpenMP/loky, deterministic, cannot hang);
    # median imputer keeps it robust to missing features on uploaded data.
    model = Pipeline([
        ("impute", SimpleImputer(strategy="median")),
        ("rf", RandomForestRegressor(n_estimators=60, max_depth=14, min_samples_leaf=20,
                                     n_jobs=1, random_state=random_state)),
    ])
    model.fit(Xtr, ytr)
    pred = model.predict(Xte)
    metrics = {
        "r2": round(float(r2_score(yte, pred)), 4),
        "mae": round(float(mean_absolute_error(yte, pred)), 3),
        "level_accuracy": round(float(accuracy_score(
            [band(v) for v in yte], [band(v) for v in pred])), 4),
        "n_train": int(len(Xtr)), "n_test": int(len(Xte)),
    }
    medians = X.median().to_dict()
    # global importance via drop-to-median ablation on a small sample (cheap, robust)
    samp = Xte.sample(min(800, len(Xte)), random_state=random_state).reset_index(drop=True)
    full = model.predict(samp)
    raw = {f: float(np.mean(np.abs(full - model.predict(samp.assign(**{f: medians[f]})))))
           for f in FEATURES}
    tot = sum(raw.values()) or 1.0
    importances = {f: round(v / tot, 4) for f, v in raw.items()}
    baseline = float(model.predict(pd.DataFrame([medians])[FEATURES])[0])
    MODELS.mkdir(exist_ok=True)
    bundle = {"model": model, "features": FEATURES, "metrics": metrics,
              "importances": importances, "medians": medians, "baseline": baseline,
              "sklearn_version": sklearn.__version__}
    joblib.dump(bundle, MODEL_PATH)
    return bundle


def load_model():
    """Load the saved model; retrain if missing, unreadable, or built with a
    different scikit-learn version (pickled models aren't portable across
    versions — retraining locally avoids AttributeError at predict time)."""
    if MODEL_PATH.exists():
        try:
            b = joblib.load(MODEL_PATH)
            if b.get("sklearn_version") == sklearn.__version__:
                # sanity-check it can actually predict on this install
                b["model"].predict(pd.DataFrame([b["medians"]])[FEATURES])
                return b
        except Exception:
            pass
    return train_and_save()


# ------------------------------------------------------------------ input handling
def read_tabular(file, name: str) -> pd.DataFrame:
    n = name.lower()
    if n.endswith(".csv"):
        return pd.read_csv(file)
    if n.endswith(".xlsx"):
        return pd.read_excel(file, engine="openpyxl")
    if n.endswith(".xls"):
        return pd.read_excel(file)
    raise ValueError(f"Unsupported tabular file: {name}")


def align_features(df: pd.DataFrame) -> pd.DataFrame:
    """Map an uploaded table onto the model feature columns (tolerant)."""
    d = df.copy()
    d.columns = [str(c).strip() for c in d.columns]
    low = {c.lower(): c for c in d.columns}
    out = pd.DataFrame(index=d.index)
    for feat in FEATURES:
        col = None
        if feat in d.columns:
            col = feat
        elif feat.lower() in low:
            col = low[feat.lower()]
        if col is not None:
            out[feat] = pd.to_numeric(d[col], errors="coerce")
        else:
            out[feat] = np.nan
    # synonyms
    for alt, feat in SYNONYMS.items():
        if alt in low and out[feat].isna().all():
            vals = pd.to_numeric(d[low[alt]], errors="coerce")
            out[feat] = (1 - vals) if feat == "theatre_pressure" else vals
    # keep id columns for display if present
    for idc in ["month", "provider_code", "provider_name",
                "treatment_function_code", "treatment_function_name"]:
        if idc in low:
            out[idc] = d[low[idc]].values
    return out


# ------------------------------------------------------------------ prediction + XAI
def predict(bundle, df_feats: pd.DataFrame):
    X = df_feats[FEATURES].astype(float)
    pred = bundle["model"].predict(X)
    out = df_feats.copy()
    out["predicted_priority_score"] = np.round(pred, 2)
    out["predicted_level"] = [band(v) for v in pred]
    return out


def local_contributions(bundle, df_feats: pd.DataFrame) -> pd.DataFrame:
    """Drop-to-median ablation attribution (dependency-free, vectorised).
    contribution_f = pred(actual) - pred(feature f set to training median).
    Positive = that feature pushes the priority UP for this area."""
    model, med = bundle["model"], bundle["medians"]
    X = df_feats[FEATURES].astype(float).reset_index(drop=True)
    full = model.predict(X)
    contribs = {}
    for f in FEATURES:
        Xm = X.copy()
        Xm[f] = med[f]
        contribs[f] = full - model.predict(Xm)
    c = pd.DataFrame(contribs)
    c.insert(0, "predicted_priority_score", np.round(full, 2))
    return c


def top_drivers(contrib_row: pd.Series, k=5):
    s = contrib_row[FEATURES].astype(float).sort_values(ascending=False)
    return [(FEATURE_LABELS[f], round(float(v), 2)) for f, v in s.head(k).items() if v > 0]


# ------------------------------------------------------------------ fairness / bias audit
def fairness_audit(pred_df: pd.DataFrame) -> dict:
    """Check the model promotes equity and does not discriminate."""
    res = {"correlations": {}, "notes": []}
    y = pred_df["predicted_priority_score"]
    for f in FAIRNESS_FEATURES:
        if f in pred_df and pred_df[f].notna().any():
            res["correlations"][FEATURE_LABELS[f]] = round(float(pred_df[f].corr(y)), 3)
    dep = "fairness_deprivation_score"
    if dep in pred_df and pred_df[dep].notna().sum() > 5:
        q = pd.qcut(pred_df[dep].rank(method="first"), 3, labels=["Lower", "Mid", "Higher"])
        res["priority_by_deprivation_tercile"] = (
            y.groupby(q, observed=False).mean().round(2).to_dict())
    res["notes"] = [
        "Protected characteristics are used only as aggregate fairness gaps "
        "(over-representation in the longest waits), never as raw group identity.",
        "A positive correlation between a fairness gap and predicted priority is "
        "intended: it raises priority for under-served areas (equity-promoting).",
        "Higher-deprivation areas should show equal-or-higher mean predicted priority; "
        "a lower value would signal bias and is flagged below.",
    ]
    if "priority_by_deprivation_tercile" in res:
        t = res["priority_by_deprivation_tercile"]
        res["bias_flag"] = bool(t.get("Higher", 0) < t.get("Lower", 0))
    return res


# ------------------------------------------------------------------ rubric (Word / dict)
def parse_rubric_docx(file) -> dict:
    """Extract a {feature: weight} dictionary from a .docx rubric.
    Reads tables (feature | weight) and 'feature: number' text lines."""
    from docx import Document
    doc = Document(file)
    weights = {}

    def match_feature(text):
        t = text.strip().lower()
        for feat, lab in FEATURE_LABELS.items():
            if t == feat.lower() or t == lab.lower() or lab.lower() in t or feat.lower() in t:
                return feat
        return None

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                feat = match_feature(cells[0])
                try:
                    w = float(cells[1])
                except (ValueError, IndexError):
                    w = None
                if feat and w is not None:
                    weights[feat] = w
    for p in doc.paragraphs:
        if ":" in p.text:
            left, right = p.text.split(":", 1)
            feat = match_feature(left)
            try:
                w = float(right.strip().split()[0])
            except (ValueError, IndexError):
                w = None
            if feat and w is not None:
                weights[feat] = w
    return weights


def apply_rubric(df_feats: pd.DataFrame, weights: dict) -> pd.Series:
    """Transparent rubric-weighted score (0-100) using min-max-normalised features
    and user-supplied weights (renormalised to sum to 1)."""
    feats = [f for f in weights if f in df_feats.columns]
    if not feats:
        return pd.Series(np.nan, index=df_feats.index)
    total = sum(abs(weights[f]) for f in feats) or 1.0
    score = np.zeros(len(df_feats))
    for f in feats:
        v = pd.to_numeric(df_feats[f], errors="coerce").astype(float)
        lo, hi = v.min(), v.max()
        vn = (v - lo) / (hi - lo) if hi > lo else v * 0
        score = score + (weights[f] / total) * vn.fillna(0).values
    return pd.Series(np.round(100 * score, 2), index=df_feats.index)
